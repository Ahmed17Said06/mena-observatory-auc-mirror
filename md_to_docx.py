from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

doc = Document()

# Page margins
section = doc.sections[0]
section.left_margin = Inches(1)
section.right_margin = Inches(1)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    if level == 1:
        p.runs[0].font.color.rgb = RGBColor(0x1F, 0x39, 0x6B)
    elif level == 2:
        p.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p

def add_table_from_rows(doc, rows):
    if not rows:
        return
    header = rows[0]
    data = rows[1:]

    table = doc.add_table(rows=1, cols=len(header))
    table.style = 'Table Grid'

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header):
        hdr_cells[i].text = h.strip()
        set_cell_bg(hdr_cells[i], '2E74B5')
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    col_count = len(header)
    for j, row in enumerate(data):
        # Pad or truncate row to match header column count
        padded = (row + [''] * col_count)[:col_count]
        row_cells = table.add_row().cells
        bg = 'F2F7FD' if j % 2 == 0 else 'FFFFFF'
        for i, cell_text in enumerate(padded):
            row_cells[i].text = cell_text.strip()
            set_cell_bg(row_cells[i], bg)
            run = row_cells[i].paragraphs[0].runs[0] if row_cells[i].paragraphs[0].runs else row_cells[i].paragraphs[0].add_run(cell_text.strip())
            run.font.size = Pt(10)

    # Column widths: Key=1.5", English=2.5", Arabic=2.5"
    if len(header) == 3:
        widths = [Inches(1.5), Inches(2.5), Inches(2.5)]
    elif len(header) == 4:
        widths = [Inches(1.2), Inches(1.8), Inches(2.0), Inches(1.5)]
    else:
        widths = None
    if widths:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                cell.width = widths[i]

    doc.add_paragraph()

# ── Parse the markdown ──────────────────────────────────────────────────────
with open('translation-strings.md', encoding='utf-8') as f:
    lines = f.readlines()

# Title + subtitle
add_heading(doc, 'MENA Observatory — Translation Strings', 1)
p = doc.add_paragraph(
    'For translation team use.  '
    'Column A = translation key (do not translate).  '
    'Column B = English text (source).  '
    'Column C = Arabic translation (fill in blank cells).'
)
p.runs[0].italic = True
doc.add_paragraph()

current_table_rows = []
i = 0

def flush_table(doc, rows):
    if len(rows) >= 2:
        add_table_from_rows(doc, rows)

while i < len(lines):
    line = lines[i].rstrip('\n')

    # H1
    if line.startswith('# ') and not line.startswith('## '):
        flush_table(doc, current_table_rows); current_table_rows = []
        i += 1; continue  # skip — already added as title

    # H2
    if line.startswith('## '):
        flush_table(doc, current_table_rows); current_table_rows = []
        text = line.lstrip('#').strip().lstrip('⚠️').strip()
        add_heading(doc, text, 2)
        i += 1; continue

    # Table row
    if line.startswith('|'):
        # skip separator rows (---|---|---)
        if re.match(r'^\|[\s\-|]+\|$', line):
            i += 1; continue
        cells = [c for c in line.split('|') if c != '']
        # Clean markdown escapes
        cells = [c.replace('\\|', '|').replace('\\', '').strip() for c in cells]
        current_table_rows.append(cells)
        i += 1; continue

    # Bullet / note paragraph
    if line.startswith('- '):
        flush_table(doc, current_table_rows); current_table_rows = []
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(line[2:].strip()).font.size = Pt(10)
        i += 1; continue

    # Horizontal rule
    if line.startswith('---'):
        flush_table(doc, current_table_rows); current_table_rows = []
        i += 1; continue

    # Blank line
    if line.strip() == '':
        flush_table(doc, current_table_rows); current_table_rows = []
        i += 1; continue

    # Plain paragraph (bold intro etc.)
    if line.strip():
        flush_table(doc, current_table_rows); current_table_rows = []
        doc.add_paragraph(line.strip()).runs[0].font.size = Pt(10) if doc.paragraphs[-1].runs else None
        i += 1; continue

    i += 1

flush_table(doc, current_table_rows)

out = 'translation-strings.docx'
doc.save(out)
print(f'Saved: {out}')
